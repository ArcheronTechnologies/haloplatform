"""
Document upload adapter for ingesting files for NLP analysis.

Supports:
- PDF documents
- Word documents (.docx)
- Plain text files
- Email files (.eml, .msg)

Extracts text content for the NLP pipeline.
"""

import email
import logging
import mimetypes
import re
from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional
from uuid import uuid4

from halo.ingestion.base_adapter import BaseAdapter, IngestionRecord

logger = logging.getLogger(__name__)


@dataclass
class ExtractedDocument:
    """A document with extracted text content."""

    document_id: str
    title: str
    content: str
    document_type: str
    language: str = "sv"

    # Metadata
    author: Optional[str] = None
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    page_count: int = 0

    # File info
    filename: str = ""
    file_size: int = 0
    mime_type: str = ""

    # Additional metadata
    metadata: dict = field(default_factory=dict)


class DocumentUploadAdapter(BaseAdapter):
    """
    Adapter for uploading and processing documents.

    Extracts text from various document formats for NLP analysis.
    """

    # Supported MIME types
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "text/plain": "txt",
        "text/html": "html",
        "message/rfc822": "eml",
        "application/vnd.ms-outlook": "msg",
    }

    def __init__(
        self,
        max_file_size_mb: int = 50,
        extract_images: bool = False,
    ):
        """
        Initialize the document upload adapter.

        Args:
            max_file_size_mb: Maximum file size in megabytes
            extract_images: Whether to extract text from images (OCR)
        """
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self.extract_images = extract_images

    @property
    def source_name(self) -> str:
        return "document_upload"

    async def fetch_company(self, orgnr: str) -> Optional[IngestionRecord]:
        """Not applicable for document adapter."""
        return None

    async def fetch_person(self, personnummer: str) -> Optional[IngestionRecord]:
        """Not applicable for document adapter."""
        return None

    async def search(
        self, query: str, limit: int = 10
    ):
        """
        Search for documents matching query.

        Note: Document search is handled by Elasticsearch, not this adapter.
        This method is required by the base class but not used for documents.
        """
        # Document search is done via Elasticsearch, not this adapter
        return
        yield  # Make it an async generator

    async def process_file(
        self,
        content: bytes,
        filename: str,
        mime_type: Optional[str] = None,
    ) -> IngestionRecord:
        """
        Process an uploaded file and extract text content.

        Args:
            content: File content as bytes
            filename: Original filename
            mime_type: MIME type (auto-detected if not provided)

        Returns:
            IngestionRecord with extracted text

        Raises:
            ValueError: If file type is not supported or file is too large
        """
        # Check file size
        if len(content) > self.max_file_size:
            raise ValueError(
                f"File too large: {len(content)} bytes (max {self.max_file_size})"
            )

        # Detect MIME type
        if not mime_type:
            mime_type, _ = mimetypes.guess_type(filename)

        if not mime_type or mime_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {mime_type}")

        doc_type = self.SUPPORTED_TYPES[mime_type]

        # Extract text based on type
        if doc_type == "pdf":
            extracted = await self._extract_pdf(content, filename)
        elif doc_type in ("docx", "doc"):
            extracted = await self._extract_docx(content, filename)
        elif doc_type == "txt":
            extracted = await self._extract_text(content, filename)
        elif doc_type == "html":
            extracted = await self._extract_html(content, filename)
        elif doc_type in ("eml", "msg"):
            extracted = await self._extract_email(content, filename, doc_type)
        else:
            raise ValueError(f"Extraction not implemented for: {doc_type}")

        return IngestionRecord(
            source=self.source_name,
            source_id=extracted.document_id,
            entity_type="document",
            raw_data={
                "title": extracted.title,
                "content": extracted.content,
                "document_type": extracted.document_type,
                "language": extracted.language,
                "author": extracted.author,
                "created_date": extracted.created_date.isoformat()
                if extracted.created_date
                else None,
                "page_count": extracted.page_count,
                "filename": extracted.filename,
                "file_size": extracted.file_size,
                "mime_type": extracted.mime_type,
                "metadata": extracted.metadata,
            },
            fetched_at=datetime.utcnow(),
        )

    async def _extract_pdf(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """
        Extract text from PDF file.

        Uses PyMuPDF (fitz) for extraction.
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF (fitz) is required for PDF extraction")

        doc = fitz.open(stream=content, filetype="pdf")

        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())

        full_text = "\n\n".join(text_parts)

        # Extract metadata
        metadata = doc.metadata or {}

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=metadata.get("title", filename),
            content=full_text,
            document_type="pdf",
            author=metadata.get("author"),
            created_date=self._parse_pdf_date(metadata.get("creationDate")),
            modified_date=self._parse_pdf_date(metadata.get("modDate")),
            page_count=len(doc),
            filename=filename,
            file_size=len(content),
            mime_type="application/pdf",
            metadata=metadata,
        )

    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date format (D:YYYYMMDDHHmmSS)."""
        if not date_str:
            return None

        try:
            # Remove D: prefix and timezone
            date_str = date_str.replace("D:", "")[:14]
            return datetime.strptime(date_str, "%Y%m%d%H%M%S")
        except (ValueError, IndexError):
            return None

    async def _extract_docx(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """
        Extract text from Word document.

        Uses python-docx for extraction.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word document extraction")

        doc = Document(BytesIO(content))

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        # Extract metadata
        core_props = doc.core_properties

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=core_props.title or filename,
            content=full_text,
            document_type="docx",
            author=core_props.author,
            created_date=core_props.created,
            modified_date=core_props.modified,
            page_count=0,  # Not easily available in docx
            filename=filename,
            file_size=len(content),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
            },
        )

    async def _extract_text(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """Extract text from plain text file."""
        # Try common encodings
        encodings = ["utf-8", "iso-8859-1", "cp1252"]

        text = None
        used_encoding = "utf-8"

        for encoding in encodings:
            try:
                text = content.decode(encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            # Fallback with replacement
            text = content.decode("utf-8", errors="replace")

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=filename,
            content=text,
            document_type="txt",
            filename=filename,
            file_size=len(content),
            mime_type="text/plain",
            metadata={"encoding": used_encoding},
        )

    async def _extract_html(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """
        Extract text from HTML file.

        Uses BeautifulSoup for extraction.
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required for HTML extraction")

        # Decode content
        text_content = content.decode("utf-8", errors="replace")

        soup = BeautifulSoup(text_content, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "head"]):
            element.decompose()

        # Get text
        text = soup.get_text(separator="\n")

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        text = "\n".join(line for line in lines if line)

        # Extract title
        title = filename
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text().strip()

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=title,
            content=text,
            document_type="html",
            filename=filename,
            file_size=len(content),
            mime_type="text/html",
        )

    async def _extract_email(
        self,
        content: bytes,
        filename: str,
        email_type: str,
    ) -> ExtractedDocument:
        """
        Extract text from email file.

        Handles both .eml and .msg formats.
        """
        if email_type == "eml":
            return await self._extract_eml(content, filename)
        else:
            return await self._extract_msg(content, filename)

    async def _extract_eml(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """Extract text from .eml file."""
        msg = email.message_from_bytes(content)

        # Extract headers
        subject = msg.get("Subject", filename)
        from_addr = msg.get("From", "")
        to_addr = msg.get("To", "")
        date_str = msg.get("Date", "")

        # Parse date
        email_date = None
        if date_str:
            try:
                email_date = email.utils.parsedate_to_datetime(date_str)
            except (ValueError, TypeError):
                pass

        # Extract body
        text_parts = []

        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        text_parts.append(payload.decode(charset, errors="replace"))
                elif content_type == "text/html":
                    # Extract text from HTML if no plain text
                    payload = part.get_payload(decode=True)
                    if payload and not text_parts:
                        try:
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(payload, "html.parser")
                            text_parts.append(soup.get_text(separator="\n"))
                        except ImportError:
                            pass
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or "utf-8"
                text_parts.append(payload.decode(charset, errors="replace"))

        full_text = "\n\n".join(text_parts)

        # Prepend email headers
        header_text = f"From: {from_addr}\nTo: {to_addr}\nSubject: {subject}\nDate: {date_str}\n\n"
        full_text = header_text + full_text

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=subject,
            content=full_text,
            document_type="email",
            author=from_addr,
            created_date=email_date,
            filename=filename,
            file_size=len(content),
            mime_type="message/rfc822",
            metadata={
                "from": from_addr,
                "to": to_addr,
                "subject": subject,
            },
        )

    async def _extract_msg(
        self,
        content: bytes,
        filename: str,
    ) -> ExtractedDocument:
        """
        Extract text from Outlook .msg file.

        Requires extract-msg library.
        """
        try:
            import extract_msg
        except ImportError:
            raise ImportError("extract-msg is required for .msg file extraction")

        msg = extract_msg.Message(BytesIO(content))

        subject = msg.subject or filename
        from_addr = msg.sender or ""
        to_addr = msg.to or ""
        body = msg.body or ""

        # Parse date
        email_date = None
        if msg.date:
            email_date = msg.date

        # Prepend headers
        header_text = f"From: {from_addr}\nTo: {to_addr}\nSubject: {subject}\n\n"
        full_text = header_text + body

        return ExtractedDocument(
            document_id=str(uuid4()),
            title=subject,
            content=full_text,
            document_type="email",
            author=from_addr,
            created_date=email_date,
            filename=filename,
            file_size=len(content),
            mime_type="application/vnd.ms-outlook",
            metadata={
                "from": from_addr,
                "to": to_addr,
                "subject": subject,
            },
        )

    def detect_language(self, text: str) -> str:
        """
        Detect the language of text.

        Simple heuristic for Swedish vs English.
        """
        # Swedish-specific characters and words
        swedish_indicators = [
            "å", "ä", "ö",  # Swedish chars
            " och ", " att ", " med ", " för ", " den ", " det ",
            " är ", " har ", " som ", " till ", " av ",
        ]

        english_indicators = [
            " the ", " and ", " is ", " for ", " with ",
            " that ", " have ", " from ", " this ",
        ]

        text_lower = text.lower()

        swedish_score = sum(1 for ind in swedish_indicators if ind in text_lower)
        english_score = sum(1 for ind in english_indicators if ind in text_lower)

        if swedish_score > english_score:
            return "sv"
        elif english_score > swedish_score:
            return "en"
        else:
            return "sv"  # Default to Swedish

    async def healthcheck(self) -> bool:
        """Check if required libraries are available."""
        try:
            import fitz  # noqa: F401
            return True
        except ImportError:
            logger.warning("PyMuPDF not available for PDF extraction")
            return True  # Still functional for other formats

    async def close(self) -> None:
        """No cleanup needed."""
        pass
